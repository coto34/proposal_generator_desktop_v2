# services/document_processor.py - Enhanced with better error handling and format support
import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
import json
import tempfile
from datetime import datetime

# Enhanced imports with better error handling
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logging.warning("PyPDF not available for PDF processing")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available for DOCX processing")

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False
    logging.warning("pdfminer not available for PDF processing")

try:
    from docxtpl import DocxTemplate
    DOCXTPL_AVAILABLE = True
except ImportError:
    DOCXTPL_AVAILABLE = False
    logging.warning("docxtpl not available for template processing")

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logging.warning("openpyxl not available for Excel processing")


# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessorError(Exception):
    """Custom exception for document processing errors"""
    pass


class DocumentProcessor:
    """Enhanced document processor with comprehensive error handling and format support"""
    
    SUPPORTED_FORMATS = {
        '.pdf': 'PDF Document',
        '.docx': 'Microsoft Word Document',
        '.doc': 'Legacy Microsoft Word Document',
        '.txt': 'Plain Text Document',
        '.rtf': 'Rich Text Format Document'
    }
    
    @staticmethod
    def check_dependencies() -> Dict[str, bool]:
        """Check availability of required dependencies"""
        return {
            'pypdf': PYPDF_AVAILABLE,
            'python-docx': DOCX_AVAILABLE,
            'pdfminer': PDFMINER_AVAILABLE,
            'docxtpl': DOCXTPL_AVAILABLE,
            'openpyxl': OPENPYXL_AVAILABLE
        }
    @staticmethod
    def chunk_document(self, content: str) -> dict:
        """
        Chunk document content for processing
        Returns a dictionary with sections and paragraphs
        """
        print(f"DEBUG: Starting chunk_document with {len(content)} characters")
        
        if not content:
            return {"sections": [], "paragraphs": []}
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        print(f"DEBUG: Found {len(paragraphs)} paragraphs")
        
        # Try to identify sections based on common patterns
        sections = []
        current_section = ""
        section_patterns = [
            r'(?i)^\s*(?:\d+\.?\s*)?(?:background|antecedentes|contexto)',
            r'(?i)^\s*(?:\d+\.?\s*)?(?:objective|objetivo|propósito)',
            r'(?i)^\s*(?:\d+\.?\s*)?(?:scope|alcance|ámbito)',
            r'(?i)^\s*(?:\d+\.?\s*)?(?:methodology|metodología|enfoque)',
            r'(?i)^\s*(?:\d+\.?\s*)?(?:activities|actividades|acciones)',
            r'(?i)^\s*(?:\d+\.?\s*)?(?:deliverable|entregable|producto)',
            r'(?i)^\s*(?:\d+\.?\s*)?(?:timeline|cronograma|calendario)',
            r'(?i)^\s*(?:\d+\.?\s*)?(?:budget|presupuesto|costo)',
        ]
        
        import re
        
        for paragraph in paragraphs:
            is_section_header = False
            for pattern in section_patterns:
                if re.match(pattern, paragraph):
                    if current_section:
                        sections.append(current_section)
                    current_section = paragraph
                    is_section_header = True
                    break
            
            if not is_section_header and current_section:
                current_section += "\n\n" + paragraph
        
        # Add the last section
        if current_section:
            sections.append(current_section)
        
        # If no sections found, treat the whole document as one section
        if not sections:
            sections = [content]
        
        print(f"DEBUG: Identified {len(sections)} sections")
        
        result = {
            "sections": sections,
            "paragraphs": paragraphs,
            "section_count": len(sections),
            "paragraph_count": len(paragraphs),
            "word_count": len(content.split()),
            "char_count": len(content)
        }
        
        print("DEBUG: chunk_document completed successfully")
        return result
    @staticmethod
    def get_missing_dependencies() -> List[str]:
        """Get list of missing dependencies with install commands"""
        missing = []
        deps = DocumentProcessor.check_dependencies()
        
        if not deps['pypdf']:
            missing.append("pip install pypdf")
        if not deps['python-docx']:
            missing.append("pip install python-docx")
        if not deps['pdfminer']:
            missing.append("pip install pdfminer.six")
        if not deps['docxtpl']:
            missing.append("pip install docxtpl")
        if not deps['openpyxl']:
            missing.append("pip install openpyxl")
        
        return missing
    @staticmethod
    def chunk_document(content: str) -> Dict[str, Any]:
        """Simple document chunking"""
        if not content or content.startswith("Error"):
            return {'sections': [], 'paragraphs': []}
        
        # Split into sections (double newlines)
        sections = [s.strip() for s in content.split('\n\n') if s.strip()]
        
        # Split into paragraphs (single newlines)  
        paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
        
        return {
            'sections': sections,
            'paragraphs': paragraphs,
            'total_length': len(content),
            'section_count': len(sections),
            'paragraph_count': len(paragraphs)
        }
    
    @staticmethod
    def validate_file(file_path: str) -> Tuple[bool, str]:
        """Validate file exists and is supported format"""
        if not file_path:
            return False, "No file path provided"
        
        path_obj = Path(file_path)
        
        if not path_obj.exists():
            return False, f"File does not exist: {file_path}"
        
        if not path_obj.is_file():
            return False, f"Path is not a file: {file_path}"
        
        file_size = path_obj.stat().st_size
        if file_size == 0:
            return False, "File is empty"
        
        # Check file size (max 50MB)
        max_size = 50 * 1024 * 1024  # 50MB
        if file_size > max_size:
            return False, f"File too large: {file_size / (1024*1024):.1f}MB (max: 50MB)"
        
        extension = path_obj.suffix.lower()
        if extension not in DocumentProcessor.SUPPORTED_FORMATS:
            supported = ", ".join(DocumentProcessor.SUPPORTED_FORMATS.keys())
            return False, f"Unsupported format: {extension}. Supported: {supported}"
        
        return True, "File is valid"
    
    @staticmethod
    def extract_text_from_file(file_path: str, max_retries: int = 3) -> Optional[str]:
        """
        Enhanced text extraction with multiple strategies and retry logic
        """
        # Validate file first
        is_valid, error_msg = DocumentProcessor.validate_file(file_path)
        if not is_valid:
            return f"Error: {error_msg}"
        
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        logger.info(f"Extracting text from {file_path.name} ({extension})")
        
        # Route to appropriate extraction method
        extraction_methods = {
            '.pdf': DocumentProcessor._extract_from_pdf_enhanced,
            '.docx': DocumentProcessor._extract_from_docx_enhanced,
            '.doc': DocumentProcessor._extract_from_doc,
            '.txt': DocumentProcessor._extract_from_txt,
            '.rtf': DocumentProcessor._extract_from_rtf
        }
        
        extraction_method = extraction_methods.get(extension)
        if not extraction_method:
            return f"Error: Unsupported file format: {extension}"
        
        # Try extraction with retries
        for attempt in range(max_retries):
            try:
                logger.info(f"Extraction attempt {attempt + 1}/{max_retries}")
                result = extraction_method(file_path)
                
                if result and not result.startswith("Error"):
                    # Validate extracted content
                    if DocumentProcessor._validate_extracted_content(result):
                        logger.info(f"Successfully extracted {len(result)} characters")
                        return DocumentProcessor._clean_extracted_text(result)
                    else:
                        logger.warning("Extracted content failed validation")
                
            except Exception as e:
                logger.warning(f"Attempt {attempt + 1} failed: {str(e)}")
                if attempt == max_retries - 1:
                    return f"Error: Failed after {max_retries} attempts: {str(e)}"
        
        return "Error: Could not extract text from file"
    
    @staticmethod
    def _extract_from_pdf_enhanced(file_path: Path) -> str:
        """Enhanced PDF extraction with multiple strategies"""
        strategies = []
        
        # Strategy 1: pdfminer (most accurate for complex layouts)
        if PDFMINER_AVAILABLE:
            strategies.append(('pdfminer', DocumentProcessor._pdf_extract_pdfminer))
        
        # Strategy 2: pypdf (faster, good for simple layouts)
        if PYPDF_AVAILABLE:
            strategies.append(('pypdf', DocumentProcessor._pdf_extract_pypdf))
        
        if not strategies:
            return "Error: No PDF processing libraries available. Install pypdf or pdfminer.six"
        
        # Try strategies in order
        for strategy_name, strategy_func in strategies:
            try:
                logger.info(f"Trying PDF extraction with {strategy_name}")
                result = strategy_func(file_path)
                
                if result and len(result.strip()) > 100:  # Minimum content threshold
                    logger.info(f"PDF extraction successful with {strategy_name}")
                    return result
                
            except Exception as e:
                logger.warning(f"PDF extraction with {strategy_name} failed: {str(e)}")
                continue
        
        return "Error: Could not extract text from PDF using any available method"
    
    @staticmethod
    def _pdf_extract_pdfminer(file_path: Path) -> str:
        """Extract text using pdfminer"""
        try:
            text = pdfminer_extract(str(file_path))
            return text.strip()
        except Exception as e:
            raise DocumentProcessorError(f"pdfminer extraction failed: {str(e)}")
    
    @staticmethod
    def _pdf_extract_pypdf(file_path: Path) -> str:
        """Extract text using pypdf"""
        try:
            reader = pypdf.PdfReader(str(file_path))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting from page {page_num + 1}: {str(e)}")
                    continue
            
            return "\n".join(text_parts)
            
        except Exception as e:
            raise DocumentProcessorError(f"pypdf extraction failed: {str(e)}")
    
    @staticmethod
    def _extract_from_docx_enhanced(file_path: Path) -> str:
        """Enhanced DOCX extraction with table and header support"""
        if not DOCX_AVAILABLE:
            return "Error: python-docx not available. Install with: pip install python-docx"
        
        try:
            doc = Document(str(file_path))
            content_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_text = DocumentProcessor._extract_table_content(table)
                if table_text:
                    content_parts.append(f"\n[TABLA]\n{table_text}\n[/TABLA]\n")
            
            # Extract headers and footers
            for section in doc.sections:
                # Headers
                if section.header:
                    for para in section.header.paragraphs:
                        text = para.text.strip()
                        if text:
                            content_parts.append(f"[ENCABEZADO] {text}")
                
                # Footers
                if section.footer:
                    for para in section.footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            content_parts.append(f"[PIE DE PÁGINA] {text}")
            
            return "\n".join(content_parts)
            
        except Exception as e:
            raise DocumentProcessorError(f"DOCX extraction failed: {str(e)}")
    
    @staticmethod
    def _extract_table_content(table) -> str:
        """Extract content from Word table"""
        table_rows = []
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_cells.append(cell_text)
            
            if any(cell.strip() for cell in row_cells):  # Skip empty rows
                table_rows.append(" | ".join(row_cells))
        
        return "\n".join(table_rows)
    
    @staticmethod
    def _extract_from_doc(file_path: Path) -> str:
        """Extract from legacy .doc format"""
        # Note: Legacy .doc format requires additional libraries like python-docx2txt or antiword
        # For now, provide helpful error message
        return ("Error: Legacy .doc format not supported. "
                "Please convert to .docx format or use a conversion tool.")
    
    @staticmethod
    def _extract_from_txt(file_path: Path) -> str:
        """Extract from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    logger.info(f"Successfully read text file with {encoding} encoding")
                    return content
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessorError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise DocumentProcessorError(f"Text file extraction failed: {str(e)}")
    
    @staticmethod
    def _extract_from_rtf(file_path: Path) -> str:
        """Extract from RTF format"""
        # RTF extraction would require additional libraries like striprtf
        return ("Error: RTF format not supported. "
                "Please convert to .docx or .txt format.")
    
    @staticmethod
    def _validate_extracted_content(content: str) -> bool:
        """Validate extracted content quality"""
        if not content or not content.strip():
            return False
        
        # Minimum length check
        if len(content.strip()) < 50:
            return False
        
        # Check for reasonable text-to-symbol ratio
        text_chars = sum(1 for c in content if c.isalnum() or c.isspace())
        total_chars = len(content)
        
        if total_chars > 0 and (text_chars / total_chars) < 0.3:
            logger.warning("Content appears to have low text-to-symbol ratio")
            return False
        
        return True
    
    @staticmethod
    def _clean_extracted_text(content: str) -> str:
        """Clean and normalize extracted text"""
        import re
        
        # Remove excessive whitespace
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)  # Multiple newlines to double
        content = re.sub(r'[ \t]+', ' ', content)  # Multiple spaces to single
        
        # Remove common PDF artifacts
        content = re.sub(r'\f', '\n', content)  # Form feed to newline
        content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', content)  # Control chars
        
        # Clean up line breaks
        content = re.sub(r'([.!?])\s*\n\s*([A-Z])', r'\1 \2', content)  # Join broken sentences
        
        return content.strip()
    
    @staticmethod
    def generate_docx_from_template(template_path: Optional[str], output_path: str, 
                                  context: Dict[str, Any]) -> bool:
        """
        Enhanced DOCX generation with better template handling
        """
        try:
            logger.info(f"Generating DOCX document: {output_path}")
            
            if template_path and os.path.exists(template_path):
                return DocumentProcessor._generate_with_template(template_path, output_path, context)
            else:
                return DocumentProcessor._generate_basic_docx(output_path, context)
                
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            return False
    
    @staticmethod
    def _generate_with_template(template_path: str, output_path: str, context: Dict[str, Any]) -> bool:
        """Generate DOCX using template"""
        if not DOCXTPL_AVAILABLE:
            logger.warning("docxtpl not available, falling back to basic generation")
            return DocumentProcessor._generate_basic_docx(output_path, context)
        
        try:
            # Validate template
            if not DocumentProcessor._validate_template(template_path):
                logger.warning("Template validation failed, using basic generation")
                return DocumentProcessor._generate_basic_docx(output_path, context)
            
            doc = DocxTemplate(template_path)
            
            # Prepare context with safe values
            safe_context = DocumentProcessor._prepare_template_context(context)
            
            doc.render(safe_context)
            doc.save(output_path)
            
            logger.info("Document generated successfully with template")
            return True
            
        except Exception as e:
            logger.error(f"Template generation failed: {str(e)}")
            logger.info("Falling back to basic generation")
            return DocumentProcessor._generate_basic_docx(output_path, context)
    
    @staticmethod
    def _validate_template(template_path: str) -> bool:
        """Validate template file"""
        try:
            if not DOCX_AVAILABLE:
                return False
            
            # Try to open template
            doc = Document(template_path)
            
            # Check if it has content
            has_content = any(para.text.strip() for para in doc.paragraphs)
            
            return has_content
            
        except Exception:
            return False
    
    @staticmethod
    def _prepare_template_context(context: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare context for template rendering"""
        safe_context = {}
        
        for key, value in context.items():
            if value is None:
                safe_context[key] = ""
            elif isinstance(value, (str, int, float, bool)):
                safe_context[key] = str(value)
            elif isinstance(value, dict):
                # Flatten dict for template use
                for sub_key, sub_value in value.items():
                    safe_key = f"{key}_{sub_key}"
                    safe_context[safe_key] = str(sub_value) if sub_value is not None else ""
            else:
                safe_context[key] = str(value)
        
        # Add current date
        safe_context['current_date'] = datetime.now().strftime("%Y-%m-%d")
        safe_context['current_datetime'] = datetime.now().strftime("%Y-%m-%d %H:%M")
        
        return safe_context
    
    @staticmethod
    def _generate_basic_docx(output_path: str, context: Dict[str, Any]) -> bool:
        """Generate basic DOCX without template"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for basic generation")
            return False
        
        try:
            doc = Document()
            
            # Document title
            title = context.get('project_title') or context.get('title', 'Propuesta de Proyecto')
            title_para = doc.add_heading(title, 0)
            
            # Project information section
            doc.add_heading('Información del Proyecto', level=1)
            
            project_info = [
                ('País', context.get('country', 'N/A')),
                ('Ubicación', f"{context.get('municipality', '')}, {context.get('department', '')}".strip(', ')),
                ('Donante', context.get('donor', 'N/A')),
                ('Duración', f"{context.get('duration_months', 'N/A')} meses"),
                ('Presupuesto', context.get('budget_cap', 'A determinar')),
                ('Beneficiarios directos', context.get('beneficiaries_direct', 'N/A')),
                ('Enfoque demográfico', context.get('demographic_focus', 'N/A'))
            ]
            
            for label, value in project_info:
                if value and value != 'N/A':
                    para = doc.add_paragraph()
                    para.add_run(f'{label}: ').bold = True
                    para.add_run(str(value))
            
            # Organization profile
            if context.get('org_profile'):
                doc.add_heading('Perfil de la Organización Ejecutora', level=1)
                doc.add_paragraph(context.get('org_profile'))
            
            # Narrative content
            if context.get('narrative'):
                doc.add_heading('Narrativa del Proyecto', level=1)
                
                # Split narrative into sections if it contains section headers
                narrative_text = context.get('narrative', '')
                sections = DocumentProcessor._split_narrative_sections(narrative_text)
                
                for section_title, section_content in sections:
                    if section_title:
                        doc.add_heading(section_title, level=2)
                    
                    # Split content into paragraphs
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())
            
            # Add footer
            doc.add_paragraph('\n')
            footer_para = doc.add_paragraph()
            footer_para.add_run('Documento generado por: ').bold = True
            footer_para.add_run('IEPADES - Instituto de Enseñanza para el Desarrollo Sostenible')
            footer_para.add_run(f'\nFecha de generación: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            
            doc.save(output_path)
            logger.info("Basic DOCX document generated successfully")
            return True
            
        except Exception as e:
            logger.error(f"Error creating basic DOCX: {str(e)}")
            return False
    
    @staticmethod
    def _split_narrative_sections(narrative: str) -> List[Tuple[str, str]]:
        """Split narrative into sections based on headers"""
        import re
        
        # Look for section headers (lines that start with numbers or are all caps)
        section_pattern = r'^(?:\d+\.?\s*)?([A-ZÁÉÍÓÚÑÜ\s]{10,}?)(?:\n|$)'
        
        sections = []
        current_section = ""
        current_title = None
        
        lines = narrative.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                current_section += '\n'
                continue
            
            # Check if this looks like a section header
            if re.match(section_pattern, line.upper()) and len(line) < 100:
                # Save previous section
                if current_title or current_section.strip():
                    sections.append((current_title, current_section.strip()))
                
                # Start new section
                current_title = line.title()
                current_section = ""
            else:
                current_section += line + '\n'
        
        # Add final section
        if current_title or current_section.strip():
            sections.append((current_title, current_section.strip()))
        
        return sections
    
    @staticmethod
    def generate_excel_budget(output_path: str, budget_data: Dict[str, Any]) -> bool:
        """
        Enhanced Excel budget generation with professional formatting
        """
        if not OPENPYXL_AVAILABLE:
            logger.error("openpyxl not available for Excel generation")
            return False
        
        try:
            logger.info(f"Generating Excel budget: {output_path}")
            
            wb = openpyxl.Workbook()
            
            # Remove default sheet and create our sheets
            wb.remove(wb.active)
            
            # Create multiple worksheets
            ws_summary = wb.create_sheet("Resumen Ejecutivo")
            ws_detail = wb.create_sheet("Presupuesto Detallado") 
            ws_category = wb.create_sheet("Por Categoría")
            ws_timeline = wb.create_sheet("Cronograma")
            
            # Generate each worksheet
            DocumentProcessor._create_budget_summary_sheet(ws_summary, budget_data)
            DocumentProcessor._create_detailed_budget_sheet(ws_detail, budget_data)
            DocumentProcessor._create_category_sheet(ws_category, budget_data)
            DocumentProcessor._create_timeline_sheet(ws_timeline, budget_data)
            
            # Set summary as active sheet
            wb.active = ws_summary
            
            wb.save(output_path)
            logger.info("Excel budget generated successfully")
            return True
            
        except Exception as e:
            logger.error(f"Error generating Excel budget: {str(e)}")
            return False
    
    @staticmethod
    def _create_budget_summary_sheet(ws, budget_data: Dict[str, Any]):
        """Create budget summary worksheet"""
        
        # Title
        ws['A1'] = 'RESUMEN EJECUTIVO DEL PRESUPUESTO'
        DocumentProcessor._apply_title_style(ws['A1'])
        ws.merge_cells('A1:F1')
        
        # Basic information
        row = 3
        info_items = [
            ('Moneda:', budget_data.get('currency', 'USD')),
            ('Fecha de Preparación:', budget_data.get('preparation_date', datetime.now().strftime('%Y-%m-%d'))),
            ('Preparado por:', budget_data.get('prepared_by', 'IEPADES')),
            ('Versión:', budget_data.get('budget_version', '1.0')),
            ('Total de Ítems:', len(budget_data.get('items', []))),
            ('Duración del Proyecto:', f"{budget_data.get('project_duration_months', 'N/A')} meses")
        ]
        
        for label, value in info_items:
            ws[f'A{row}'] = label
            ws[f'B{row}'] = value
            DocumentProcessor._apply_info_style(ws[f'A{row}'])
            row += 1
        
        # Financial totals
        row += 2
        ws[f'A{row}'] = 'TOTALES FINANCIEROS'
        DocumentProcessor._apply_header_style(ws[f'A{row}'])
        ws.merge_cells(f'A{row}:C{row}')
        row += 1
        
        totals = budget_data.get('financial_totals', {})
        total_items = [
            ('Subtotal:', totals.get('subtotal', 0)),
            ('Costos Administrativos:', totals.get('administrative_cost', 0)),
            ('Contingencias:', totals.get('contingency_amount', 0)),
            ('TOTAL GENERAL:', totals.get('total', 0))
        ]
        
        for label, amount in total_items:
            ws[f'A{row}'] = label
            ws[f'B{row}'] = amount
            DocumentProcessor._apply_currency_format(ws[f'B{row}'])
            
            if 'TOTAL' in label:
                DocumentProcessor._apply_total_style(ws[f'A{row}'])
                DocumentProcessor._apply_total_style(ws[f'B{row}'])
            
            row += 1
        
        # Summary by category
        row += 2
        ws[f'A{row}'] = 'RESUMEN POR CATEGORÍA'
        DocumentProcessor._apply_header_style(ws[f'A{row}'])
        ws.merge_cells(f'A{row}:C{row}')
        row += 1
        
        # Category headers
        ws[f'A{row}'] = 'Categoría'
        ws[f'B{row}'] = 'Monto (USD)'
        ws[f'C{row}'] = 'Porcentaje'
        DocumentProcessor._apply_header_style(ws[f'A{row}'])
        DocumentProcessor._apply_header_style(ws[f'B{row}'])
        DocumentProcessor._apply_header_style(ws[f'C{row}'])
        row += 1
        
        total_budget = totals.get('total', 1)
        for category, amount in budget_data.get('summary_by_category', {}).items():
            ws[f'A{row}'] = category
            ws[f'B{row}'] = amount
            ws[f'C{row}'] = f"{(amount/total_budget)*100:.1f}%"
            DocumentProcessor._apply_currency_format(ws[f'B{row}'])
            row += 1
        
        # Auto-adjust column widths
        DocumentProcessor._auto_adjust_columns(ws)
    
    @staticmethod
    def _create_detailed_budget_sheet(ws, budget_data: Dict[str, Any]):
        """Create detailed budget worksheet"""
        
        # Headers
        headers = [
            'Código', 'Act.', 'Categoría', 'Descripción', 'Unidad',
            'Cant.', 'Costo Unit.', 'Meses', 'Fase', 'Total', 'Justificación'
        ]
        
        # Title
        ws['A1'] = 'PRESUPUESTO DETALLADO'
        DocumentProcessor._apply_title_style(ws['A1'])
        ws.merge_cells(f'A1:{get_column_letter(len(headers))}1')
        
        # Headers row
        for col, header in enumerate(headers, 1):
            cell = ws[f'{get_column_letter(col)}3']
            cell.value = header
            DocumentProcessor._apply_header_style(cell)
        
        # Data rows
        row = 4
        for item in budget_data.get('items', []):
            data = [
                item.get('code', ''),
                item.get('activity_code', ''),
                item.get('category', ''),
                item.get('description', ''),
                item.get('unit', ''),
                item.get('qty', 0),
                item.get('unit_cost', 0),
                item.get('months', 1),
                item.get('phase', ''),
                item.get('qty', 0) * item.get('unit_cost', 0) * item.get('months', 1),
                item.get('justification', '')
            ]
            
            for col, value in enumerate(data, 1):
                cell = ws[f'{get_column_letter(col)}{row}']
                cell.value = value
                
                # Apply formatting
                if col in [6, 7, 8]:  # Quantity, unit cost, months
                    DocumentProcessor._apply_number_format(cell)
                elif col == 10:  # Total
                    DocumentProcessor._apply_currency_format(cell)
            
            row += 1
        
        # Add totals row
        row += 1
        ws[f'A{row}'] = 'TOTAL'
        total_cell = ws[f'{get_column_letter(10)}{row}']
        total_cell.value = budget_data.get('financial_totals', {}).get('total', 0)
        DocumentProcessor._apply_total_style(ws[f'A{row}'])
        DocumentProcessor._apply_total_style(total_cell)
        DocumentProcessor._apply_currency_format(total_cell)
        
        # Auto-adjust columns
        DocumentProcessor._auto_adjust_columns(ws)
    
    @staticmethod
    def _create_category_sheet(ws, budget_data: Dict[str, Any]):
        """Create category analysis worksheet"""
        
        ws['A1'] = 'ANÁLISIS POR CATEGORÍA'
        DocumentProcessor._apply_title_style(ws['A1'])
        ws.merge_cells('A1:D1')
        
        # Headers
        row = 3
        headers = ['Categoría', 'Monto (USD)', 'Porcentaje', 'Descripción']
        for col, header in enumerate(headers, 1):
            cell = ws[f'{get_column_letter(col)}{row}']
            cell.value = header
            DocumentProcessor._apply_header_style(cell)
        
        # Category data
        row += 1
        total_budget = budget_data.get('financial_totals', {}).get('total', 1)
        
        for category, amount in budget_data.get('summary_by_category', {}).items():
            percentage = (amount / total_budget) * 100
            
            ws[f'A{row}'] = category
            ws[f'B{row}'] = amount
            ws[f'C{row}'] = f"{percentage:.1f}%"
            ws[f'D{row}'] = DocumentProcessor._get_category_description(category)
            
            DocumentProcessor._apply_currency_format(ws[f'B{row}'])
            row += 1
        
        DocumentProcessor._auto_adjust_columns(ws)
    
    @staticmethod
    def _create_timeline_sheet(ws, budget_data: Dict[str, Any]):
        """Create timeline/cash flow worksheet"""
        
        ws['A1'] = 'CRONOGRAMA Y FLUJO DE CAJA'
        DocumentProcessor._apply_title_style(ws['A1'])
        ws.merge_cells('A1:M1')
        
        # Monthly distribution if available
        monthly_dist = budget_data.get('timeline_breakdown', {}).get('monthly_distribution', {})
        
        if monthly_dist:
            row = 3
            ws[f'A{row}'] = 'Mes'
            ws[f'B{row}'] = 'Monto (USD)'
            ws[f'C{row}'] = 'Acumulado'
            
            for col in ['A', 'B', 'C']:
                DocumentProcessor._apply_header_style(ws[f'{col}{row}'])
            
            row += 1
            cumulative = 0
            
            for month, amount in sorted(monthly_dist.items()):
                cumulative += amount
                ws[f'A{row}'] = month
                ws[f'B{row}'] = amount
                ws[f'C{row}'] = cumulative
                
                DocumentProcessor._apply_currency_format(ws[f'B{row}'])
                DocumentProcessor._apply_currency_format(ws[f'C{row}'])
                row += 1
        
        DocumentProcessor._auto_adjust_columns(ws)
    
    @staticmethod
    def _get_category_description(category: str) -> str:
        """Get description for budget category"""
        descriptions = {
            'Personal Técnico y Profesional': 'Salarios y honorarios del equipo técnico',
            'Equipamiento, Suministros y Materiales': 'Equipos, materiales y suministros para el proyecto',
            'Viajes, Transporte y Logística': 'Costos de movilización y logística',
            'Capacitación, Talleres y Eventos': 'Actividades de formación y eventos',
            'Gastos Operativos y Servicios': 'Gastos operativos generales',
            'Costos Administrativos': 'Gastos administrativos de la organización',
            'Gastos Financieros y Bancarios': 'Comisiones bancarias y gastos financieros',
            'Contingencias': 'Reserva para imprevistos'
        }
        return descriptions.get(category, 'Descripción no disponible')
    
    # Styling methods
    @staticmethod
    def _apply_title_style(cell):
        """Apply title style to cell"""
        cell.font = Font(bold=True, size=16, color='FFFFFF')
        cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    @staticmethod
    def _apply_header_style(cell):
        """Apply header style to cell"""
        cell.font = Font(bold=True, size=12)
        cell.fill = PatternFill(start_color='D9E2F3', end_color='D9E2F3', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
    
    @staticmethod
    def _apply_total_style(cell):
        """Apply total/summary style to cell"""
        cell.font = Font(bold=True, size=12)
        cell.fill = PatternFill(start_color='FFE699', end_color='FFE699', fill_type='solid')
    
    @staticmethod
    def _apply_info_style(cell):
        """Apply info label style to cell"""
        cell.font = Font(bold=True)
    
    @staticmethod
    def _apply_currency_format(cell):
        """Apply currency formatting to cell"""
        cell.number_format = '"$"#,##0.00'
    
    @staticmethod
    def _apply_number_format(cell):
        """Apply number formatting to cell"""
        cell.number_format = '#,##0.00'
    
    @staticmethod
    def _auto_adjust_columns(ws):
        """Auto-adjust column widths"""
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            adjusted_width = min(max_length + 2, 50)  # Max width of 50
            ws.column_dimensions[column_letter].width = adjusted_width
    
    @staticmethod
    def create_processing_report(file_path: str, extraction_result: str, 
                               output_dir: Optional[str] = None) -> str:
        """Create a detailed processing report"""
        
        if not output_dir:
            output_dir = Path(file_path).parent
        
        report_path = Path(output_dir) / f"processing_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        # Analyze extracted content
        content_stats = DocumentProcessor._analyze_content(extraction_result)
        
        # File information
        file_info = DocumentProcessor._get_file_info(file_path)
        
        report_data = {
            "processing_metadata": {
                "timestamp": datetime.now().isoformat(),
                "processor_version": "2.0.0",
                "file_path": str(file_path)
            },
            "file_information": file_info,
            "extraction_results": {
                "success": not extraction_result.startswith("Error"),
                "content_length": len(extraction_result),
                "error_message": extraction_result if extraction_result.startswith("Error") else None
            },
            "content_analysis": content_stats,
            "dependencies_status": DocumentProcessor.check_dependencies(),
            "recommendations": DocumentProcessor._generate_recommendations(extraction_result, file_info)
        }
        
        try:
            with open(report_path, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
            
            logger.info(f"Processing report created: {report_path}")
            return str(report_path)
            
        except Exception as e:
            logger.error(f"Could not create processing report: {str(e)}")
            return ""
    
    @staticmethod
    def _analyze_content(content: str) -> Dict[str, Any]:
        """Analyze extracted content"""
        if not content or content.startswith("Error"):
            return {"error": "No content to analyze"}
        
        import re
        
        return {
            "character_count": len(content),
            "word_count": len(content.split()),
            "paragraph_count": len([p for p in content.split('\n\n') if p.strip()]),
            "line_count": len(content.split('\n')),
            "has_tables": "[TABLA]" in content,
            "has_headers": any(marker in content for marker in ["[ENCABEZADO]", "[PIE DE PÁGINA]"]),
            "language_detected": DocumentProcessor._detect_language(content),
            "readability_score": DocumentProcessor._calculate_readability(content)
        }
    
    @staticmethod
    def _get_file_info(file_path: str) -> Dict[str, Any]:
        """Get detailed file information"""
        try:
            path_obj = Path(file_path)
            stat = path_obj.stat()
            
            return {
                "filename": path_obj.name,
                "extension": path_obj.suffix,
                "size_bytes": stat.st_size,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "format_supported": path_obj.suffix.lower() in DocumentProcessor.SUPPORTED_FORMATS
            }
        except Exception as e:
            return {"error": f"Could not get file info: {str(e)}"}
    
    @staticmethod
    def _detect_language(content: str) -> str:
        """Simple language detection"""
        spanish_indicators = ['el', 'la', 'de', 'que', 'y', 'es', 'en', 'un', 'se', 'no', 'te', 'lo', 'le', 'da', 'su', 'por', 'son', 'con', 'para', 'al', 'proyecto', 'desarrollo']
        english_indicators = ['the', 'and', 'of', 'to', 'a', 'in', 'is', 'it', 'you', 'that', 'he', 'was', 'for', 'on', 'are', 'as', 'with', 'his', 'they', 'project', 'development']
        
        words = content.lower().split()[:200]  # Sample first 200 words
        
        spanish_count = sum(1 for word in words if word in spanish_indicators)
        english_count = sum(1 for word in words if word in english_indicators)
        
        if spanish_count > english_count:
            return "Spanish (likely)"
        elif english_count > spanish_count:
            return "English (likely)"
        else:
            return "Unknown"
    
    @staticmethod
    def _calculate_readability(content: str) -> Dict[str, float]:
        """Calculate basic readability metrics"""
        import re
        
        sentences = re.split(r'[.!?]+', content)
        words = content.split()
        
        if not sentences or not words:
            return {"error": "Insufficient content for analysis"}
        
        avg_sentence_length = len(words) / len(sentences)
        
        # Count syllables (simple approximation)
        def count_syllables(word):
            word = word.lower()
            vowels = 'aeiouy'
            syllable_count = 0
            previous_was_vowel = False
            
            for char in word:
                if char in vowels:
                    if not previous_was_vowel:
                        syllable_count += 1
                    previous_was_vowel = True
                else:
                    previous_was_vowel = False
            
            # Adjust for silent 'e'
            if word.endswith('e') and syllable_count > 1:
                syllable_count -= 1
            
            return max(1, syllable_count)
        
        total_syllables = sum(count_syllables(word) for word in words)
        avg_syllables_per_word = total_syllables / len(words) if words else 0
        
        # Simple readability score (lower is easier)
        readability_score = avg_sentence_length + avg_syllables_per_word
        
        return {
            "avg_sentence_length": round(avg_sentence_length, 2),
            "avg_syllables_per_word": round(avg_syllables_per_word, 2),
            "readability_score": round(readability_score, 2),
            "complexity": "Simple" if readability_score < 15 else "Moderate" if readability_score < 25 else "Complex"
        }
    
    @staticmethod
    def _generate_recommendations(extraction_result: str, file_info: Dict[str, Any]) -> List[str]:
        """Generate processing recommendations"""
        recommendations = []
        
        if extraction_result.startswith("Error"):
            recommendations.append("Consider converting the file to a supported format (.docx, .pdf, .txt)")
            recommendations.append("Check if the file is corrupted or password-protected")
        
        file_size_mb = file_info.get("size_mb", 0)
        if file_size_mb > 10:
            recommendations.append("Large file detected - consider splitting into smaller sections for better processing")
        
        if not file_info.get("format_supported", True):
            recommendations.append("File format not fully supported - conversion recommended")
        
        missing_deps = DocumentProcessor.get_missing_dependencies()
        if missing_deps:
            recommendations.append(f"Install missing dependencies for better processing: {', '.join(missing_deps)}")
        
        return recommendations