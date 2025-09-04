# ui/wizard.py - Enhanced version with international standards and professional budget generation
import os
import sys
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import json
import traceback
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, Callable, List
import asyncio

from ui.components import LabeledEntry, FilePicker
from services.llm_providers import DeepSeekClient, SonnetClient, create_test_clients
from services.document_processor import DocumentProcessor
from services.token_manager import TokenManager, EnhancedChainedPromptGenerator
from validation.schemas import BudgetResult

try:
    from services.enhanced_budget_system import InternationalBudgetGenerator
    ENHANCED_BUDGET_AVAILABLE = True
except ImportError:
    ENHANCED_BUDGET_AVAILABLE = False
    print("Enhanced budget system not available")

try:
    from services.academic_reference_system import AcademicReferenceSystem
    ACADEMIC_REFS_AVAILABLE = True
except ImportError:
    ACADEMIC_REFS_AVAILABLE = False
    print("Academic reference system not available")

try:
    from services.real_time_update_system import RealTimeDataIntegrator
    REALTIME_DATA_AVAILABLE = True
except ImportError:
    REALTIME_DATA_AVAILABLE = False
    print("Real-time data system not available")

try:
    from services.gantt_chart_generator import GanttChartGenerator
    GANTT_AVAILABLE = True
except ImportError:
    GANTT_AVAILABLE = False
    print("Gantt chart system not available")

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    from docxtpl import DocxTemplate
except Exception:
    DocxTemplate = None

try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
except Exception:
    openpyxl = None
    Workbook = None


class StateManager:
    def __init__(self):
        self._state = {
            "project": {},
            "tor_path": None,
            "tor_content": None,
            "tor_chunks": {},
            "models": {
                "narrative": "DeepSeek",
                "budget": "Sonnet",
                "temperature": 0.2,
                "max_tokens": 4000,
                "language": "es",
                "currency": "USD"
            },
            "templates": {"docx": None, "xlsx": None},
            "results": {"narrative": None, "budget": None, "output_paths": {}},
            "api_status": {"deepseek": False, "sonnet": False},
            "processing": {"active": False, "current_step": None, "progress": 0}
        }
        self._load_state()

    def get(self, key: str, default=None):
        keys = key.split('.')
        value = self._state
        for k in keys:
            if not isinstance(value, dict):
                return default
            value = value.get(k, default)
            if value is None:
                return default
        return value

    def set(self, key: str, value: Any) -> None:
        keys = key.split('.')
        target = self._state
        for k in keys[:-1]:
            target = target.setdefault(k, {})
        target[keys[-1]] = value
        self._save_state()

    def update(self, updates: Dict[str, Any]) -> None:
        for key, value in updates.items():
            self.set(key, value)

    def validate_project(self) -> tuple[bool, List[str]]:
        errors = []
        project = self.get("project", {})
        required = [("title", "Título del proyecto"),
                    ("country", "País"),
                    ("donor", "Donante/Financiador"),
                    ("duration_months", "Duración")]
        for field, label in required:
            if not str(project.get(field, "")).strip():
                errors.append(f"• {label} es requerido")
        duration = project.get("duration_months", "")
        if duration and not str(duration).strip().isdigit():
            errors.append("• Duración debe ser un número de meses")
        return len(errors) == 0, errors

    def validate_tor(self) -> tuple[bool, List[str]]:
        errors = []
        if not self.get("tor_path"):
            errors.append("• Falta seleccionar archivo de ToR")
        if not self.get("tor_content"):
            errors.append("• Falta procesar el contenido del ToR")
        if not self.get("tor_chunks"):
            errors.append("• Falta análisis del documento ToR")
        return len(errors) == 0, errors

    def validate_apis(self) -> tuple[bool, List[str]]:
        errors = []
        if not self.get("api_status.deepseek"):
            errors.append("• API de DeepSeek no configurada correctamente")
        if not self.get("api_status.sonnet"):
            errors.append("• API de Sonnet no configurada correctamente")
        return len(errors) == 0, errors

    def _save_state(self) -> None:
        try:
            state_dir = Path("runs/state")
            state_dir.mkdir(parents=True, exist_ok=True)
            safe_state = dict(self._state)
            safe_state.pop("tor_content", None)
            with open(state_dir / "wizard_state.json", 'w', encoding='utf-8') as f:
                json.dump(safe_state, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _load_state(self) -> None:
        try:
            state_file = Path("runs/state/wizard_state.json")
            if state_file.exists():
                with open(state_file, 'r', encoding='utf-8') as f:
                    saved_state = json.load(f)
                for key, value in saved_state.items():
                    if key in self._state and isinstance(self._state[key], dict) and isinstance(value, dict):
                        self._state[key].update(value)
                    else:
                        self._state[key] = value
        except Exception:
            pass


class ProgressManager:
    def __init__(self, progress_callback: Callable[[int, str], None]):
        self.progress_callback = progress_callback
        self.steps = []
        self.current_step = 0

    def set_steps(self, steps: List[str]) -> None:
        self.steps = steps
        self.current_step = 0

    def update(self, step_text: str, progress: int = None) -> None:
        if progress is None:
            progress = int((self.current_step / len(self.steps)) * 100) if self.steps else 0
        self.progress_callback(progress, step_text)

    def next_step(self, step_text: str = None) -> None:
        self.current_step += 1
        if step_text is None and self.current_step <= len(self.steps):
            step_text = self.steps[self.current_step - 1]
        progress = int((self.current_step / len(self.steps)) * 100) if self.steps else 0
        self.progress_callback(progress, step_text or f"Paso {self.current_step}")


class ProposalWizard(ttk.Frame):
    def __init__(self, master):
        super().__init__(master, padding=10)
        self.pack(fill="both", expand=True)

        self.state = StateManager()
        self.progress_manager = ProgressManager(self._update_progress_ui)

        self.budget_generator = None
        self.academic_system = None
        self.realtime_system = None
        self.gantt_generator = None

        self._processing = False
        self._current_thread = None

        self._setup_ui()
        self._setup_event_handlers()
        self.master.after(1000, self._check_api_status_async)

    def _extract_activities_from_budget(self, budget_data: Dict) -> List[Dict]:
        activities = []
        duration_months = int(budget_data.get("project_duration_months", 24))
        activity_mapping = {
            "Personal Técnico y Profesional": {
                "name": "Gestión y Coordinación del Proyecto",
                "duration": duration_months,
                "start_month": 1,
                "critical_path": True
            },
            "Equipamiento, Suministros y Materiales": {
                "name": "Adquisición de Equipos y Materiales",
                "duration": 3,
                "start_month": 1,
                "critical_path": False
            },
            "Capacitación, Talleres y Eventos": {
                "name": "Implementación de Capacitaciones",
                "duration": max(duration_months - 6, 1),
                "start_month": 4,
                "critical_path": True
            },
            "Viajes, Transporte y Logística": {
                "name": "Trabajo de Campo y Supervisión",
                "duration": max(duration_months - 3, 1),
                "start_month": 2,
                "critical_path": False
            }
        }
        code_counter = 1
        for category, amount in budget_data.get("summary_by_category", {}).items():
            if category in activity_mapping:
                a = activity_mapping[category]
                activities.append({
                    "code": f"A{code_counter}.1",
                    "name": a["name"],
                    "component": "Componente Principal",
                    "duration_months": a["duration"],
                    "start_month": a["start_month"],
                    "end_month": a["start_month"] + a["duration"] - 1,
                    "progress": 0,
                    "critical_path": a["critical_path"],
                    "budget_allocated": amount
                })
                code_counter += 1
        activities.extend([
            {"code": "M1", "name": "Inicio del Proyecto", "component": "Hito",
             "duration_months": 0, "start_month": 1, "end_month": 1,
             "progress": 0, "critical_path": True, "budget_allocated": 0},
            {"code": "M2", "name": "Evaluación de Medio Término", "component": "Hito",
             "duration_months": 0, "start_month": max(duration_months // 2, 1),
             "end_month": max(duration_months // 2, 1), "progress": 0,
             "critical_path": True, "budget_allocated": 0},
            {"code": "M3", "name": "Cierre del Proyecto", "component": "Hito",
             "duration_months": 0, "start_month": duration_months, "end_month": duration_months,
             "progress": 0, "critical_path": True, "budget_allocated": 0},
        ])
        return activities

    def _save_enhanced_results(self, narrative: str, budget_data: Dict, gantt_data: Dict = None, currency: str = "USD") -> Path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_folder = Path("runs") / f"proposal_enhanced_{timestamp}"
        output_folder.mkdir(parents=True, exist_ok=True)
        self._save_narrative_docx(output_folder, narrative)
        self._save_budget_xlsx(output_folder, budget_data, currency)
        if gantt_data and GANTT_AVAILABLE:
            self._save_gantt_xlsx(output_folder, gantt_data)
        self._save_json_summary(output_folder, narrative, budget_data, gantt_data)
        return output_folder

    def _save_narrative_docx(self, output_folder: Path, narrative: str):
        docx_path = output_folder / "propuesta_narrativa.docx"
        project = self.state.get("project", {})
        try:
            docx_template = self.state.get("templates.docx")
            if docx_template and DocxTemplate:
                template = DocxTemplate(docx_template)
                context = {
                    "title": project.get("title", "Propuesta de Proyecto"),
                    "country": project.get("country", "Guatemala"),
                    "donor": project.get("donor", "Donante Internacional"),
                    "duration_months": project.get("duration_months", "24"),
                    "narrative": narrative,
                    "organization": "IEPADES - Instituto de Estudios para el Desarrollo Sostenible",
                    "date": datetime.now().strftime("%d de %B de %Y")
                }
                template.render(context)
                template.save(docx_path)
                self._log_message("Narrativa guardada con plantilla DOCX")
            elif DocxDocument:
                doc = DocxDocument()
                title_para = doc.add_heading(project.get("title", "Propuesta de Proyecto"), 0)
                title_para.alignment = 1
                doc.add_paragraph()
                org_para = doc.add_paragraph("IEPADES - Instituto de Estudios para el Desarrollo Sostenible")
                org_para.alignment = 1
                date_para = doc.add_paragraph(f"Guatemala, {datetime.now().strftime('%B de %Y')}")
                date_para.alignment = 1
                doc.add_page_break()
                doc.add_paragraph(narrative)
                doc.save(docx_path)
                self._log_message("Narrativa guardada en DOCX básico")
            else:
                with open(docx_path.with_suffix('.txt'), 'w', encoding='utf-8') as f:
                    f.write(f"PROPUESTA DE PROYECTO\n\n{narrative}")
                self._log_message("Narrativa guardada como texto")
        except Exception as e:
            self._log_message(f"Error guardando narrativa: {e}")
            with open(docx_path.with_suffix('.txt'), 'w', encoding='utf-8') as f:
                f.write(narrative)

    def _save_budget_xlsx(self, output_folder: Path, budget_data: Dict, currency: str):
        xlsx_path = output_folder / "presupuesto_detallado.xlsx"
        try:
            if openpyxl:
                wb = Workbook()
                wb.remove(wb.active)
                ws_budget = wb.create_sheet("Presupuesto Detallado")
                self._create_budget_sheet(ws_budget, budget_data, currency)
                ws_summary = wb.create_sheet("Resumen Ejecutivo")
                self._create_summary_sheet(ws_summary, budget_data, currency)
                ws_categories = wb.create_sheet("Análisis por Categoría")
                self._create_category_analysis_sheet(ws_categories, budget_data, currency)
                wb.active = ws_budget
                wb.save(xlsx_path)
                self._log_message(f"Presupuesto guardado en XLSX con {len(budget_data.get('items', []))} líneas")
            else:
                self._save_budget_csv(output_folder, budget_data, currency)
        except Exception as e:
            self._log_message(f"Error guardando presupuesto XLSX: {e}")
            self._save_budget_csv(output_folder, budget_data, currency)

    def _create_budget_sheet(self, ws, budget_data: Dict, currency: str):
        project = self.state.get("project", {})
        ws['A1'] = f"PRESUPUESTO DETALLADO - {project.get('title', 'Proyecto')}"
        ws.merge_cells('A1:H1')
        ws['A3'] = "Información del Proyecto:"
        ws['A4'] = f"Donante: {project.get('donor', 'N/A')}"
        ws['A5'] = f"Duración: {project.get('duration_months', 'N/A')} meses"
        ws['A6'] = f"Moneda: {currency}"
        headers = ['Código', 'Categoría', 'Descripción', 'Unidad', 'Cantidad', 'Costo Unit.', 'Costo Total', 'Justificación']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=8, column=col)
            cell.value = header
            if openpyxl:
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color='CCCCCC', end_color='CCCCCC', fill_type='solid')
        row = 9
        for item in budget_data.get('items', []):
            ws.cell(row=row, column=1, value=item.get('code', ''))
            ws.cell(row=row, column=2, value=item.get('category', ''))
            ws.cell(row=row, column=3, value=item.get('description', ''))
            ws.cell(row=row, column=4, value=item.get('unit', ''))
            ws.cell(row=row, column=5, value=item.get('quantity', 0))
            ws.cell(row=row, column=6, value=item.get('unit_cost', 0))
            ws.cell(row=row, column=7, value=item.get('total_cost', 0))
            ws.cell(row=row, column=8, value=item.get('justification', ''))
            row += 1
        totals = budget_data.get('financial_totals', {})
        row += 1
        ws.cell(row=row, column=6, value="SUBTOTAL:")
        ws.cell(row=row, column=7, value=totals.get('subtotal', 0))
        row += 1
        ws.cell(row=row, column=6, value="Costos Administrativos:")
        ws.cell(row=row, column=7, value=totals.get('administrative_cost', 0))
        row += 1
        ws.cell(row=row, column=6, value="Contingencias:")
        ws.cell(row=row, column=7, value=totals.get('contingency_amount', 0))
        row += 1
        ws.cell(row=row, column=6, value="TOTAL GENERAL:")
        ws.cell(row=row, column=7, value=totals.get('total', 0))
        if openpyxl:
            for r in range(row-3, row+1):
                ws.cell(row=r, column=6).font = Font(bold=True)
                ws.cell(row=r, column=7).font = Font(bold=True)

    def _create_summary_sheet(self, ws, budget_data: Dict, currency: str):
        ws['A1'] = "RESUMEN EJECUTIVO DEL PRESUPUESTO"
        totals = budget_data.get('financial_totals', {})
        ws['A3'] = "Resumen Financiero"
        ws['A4'] = f"Moneda: {currency}"
        ws['A5'] = f"Total del Proyecto: {currency} {totals.get('total', 0):,.2f}"
        ws['A6'] = f"Número de Líneas Presupuestales: {len(budget_data.get('items', []))}"
        ws['A8'] = "Distribución por Categoría:"
        row = 9
        total = totals.get('total', 1) or 1
        for category, amount in budget_data.get('summary_by_category', {}).items():
            ws.cell(row=row, column=1, value=category)
            ws.cell(row=row, column=2, value=f"{currency} {amount:,.2f}")
            ws.cell(row=row, column=3, value=f"{(amount/total)*100:.1f}%")
            row += 1

    def _create_category_analysis_sheet(self, ws, budget_data: Dict, currency: str):
        ws['A1'] = "ANÁLISIS POR CATEGORÍA"
        headers = ['Categoría', 'Monto', 'Porcentaje', 'Número de Items']
        for col, header in enumerate(headers, 1):
            ws.cell(row=3, column=col, value=header)
        row = 4
        total_amount = budget_data.get('financial_totals', {}).get('total', 1) or 1
        for category, amount in budget_data.get('summary_by_category', {}).items():
            item_count = sum(1 for i in budget_data.get('items', []) if i.get('category') == category)
            ws.cell(row=row, column=1, value=category)
            ws.cell(row=row, column=2, value=f"{currency} {amount:,.2f}")
            ws.cell(row=row, column=3, value=f"{(amount/total_amount)*100:.1f}%")
            ws.cell(row=row, column=4, value=item_count)
            row += 1

    def _save_budget_csv(self, output_folder: Path, budget_data: Dict, currency: str):
        csv_path = output_folder / "presupuesto_detallado.csv"
        with open(csv_path, 'w', encoding='utf-8', newline='') as f:
            import csv
            writer = csv.writer(f)
            writer.writerow(['Código', 'Categoría', 'Descripción', 'Unidad', 'Cantidad', 'Costo Unitario', 'Costo Total', 'Moneda'])
            for item in budget_data.get('items', []):
                writer.writerow([
                    item.get('code', ''), item.get('category', ''), item.get('description', ''),
                    item.get('unit', ''), item.get('quantity', 0), item.get('unit_cost', 0),
                    item.get('total_cost', 0), currency
                ])
            totals = budget_data.get('financial_totals', {})
            writer.writerow(['', '', 'SUBTOTAL', '', '', '', totals.get('subtotal', 0), currency])
            writer.writerow(['', '', 'Costos Administrativos', '', '', '', totals.get('administrative_cost', 0), currency])
            writer.writerow(['', '', 'Contingencias', '', '', '', totals.get('contingency_amount', 0), currency])
            writer.writerow(['', '', 'TOTAL GENERAL', '', '', '', totals.get('total', 0), currency])
        self._log_message("Presupuesto guardado como CSV")

    def _save_gantt_xlsx(self, output_folder: Path, gantt_data: Dict):
        try:
            gantt_path = output_folder / "cronograma_gantt.xlsx"
            success = self.gantt_generator.generate_gantt_excel(str(gantt_path))
            if success:
                self._log_message("Cronograma Gantt guardado exitosamente")
            else:
                self._log_message("Error guardando cronograma Gantt")
        except Exception as e:
            self._log_message(f"Error creando Gantt: {e}")

    def _save_json_summary(self, output_folder: Path, narrative: str, budget_data: Dict, gantt_data: Dict):
        summary_path = output_folder / "resumen_proyecto.json"
        summary = {
            "project_info": self.state.get("project", {}),
            "generation_timestamp": datetime.now().isoformat(),
            "currency": budget_data.get("currency", "USD"),
            "total_budget": budget_data.get("financial_totals", {}).get("total", 0),
            "budget_items_count": len(budget_data.get("items", [])),
            "narrative_length": len(narrative),
            "systems_used": {
                "enhanced_budget": ENHANCED_BUDGET_AVAILABLE,
                "academic_references": ACADEMIC_REFS_AVAILABLE,
                "realtime_data": REALTIME_DATA_AVAILABLE,
                "gantt_chart": GANTT_AVAILABLE and gantt_data is not None
            },
            "files_generated": [
                "propuesta_narrativa.docx",
                "presupuesto_detallado.xlsx",
                "cronograma_gantt.xlsx" if gantt_data else None,
                "resumen_proyecto.json"
            ]
        }
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)

    def _log_message(self, message: str):
        print(f"[{datetime.now().strftime('%H:%M:%S')}] {message}")
        if hasattr(self, 'master'):
            try:
                self.master.after(0, self._add_log_entry, message)
            except:
                pass

    def _add_log_entry(self, message: str):
        if hasattr(self, 'log_text'):
            try:
                self.log_text.config(state="normal")
                self.log_text.insert(tk.END, f"[{datetime.now().strftime('%H:%M:%S')}] {message}\n")
                self.log_text.see(tk.END)
                self.log_text.config(state="disabled")
            except:
                pass

    def _setup_ui(self):
        style = ttk.Style()
        style.configure('Enhanced.TNotebook', tabposition='n')

        self.notebook = ttk.Notebook(self, style='Enhanced.TNotebook')
        self.notebook.pack(fill="both", expand=True)

        self.tabs = {}
        tab_configs = [
            ("project", "🏗️ Proyecto", self._build_project_tab),
            ("tor", "📋 Términos de Referencia", self._build_tor_tab),
            ("config", "⚙️ Configuración", self._build_config_tab),
            ("generate", "🚀 Generación", self._build_generate_tab),
            ("results", "📊 Resultados", self._build_results_tab)
        ]
        for tab_id, tab_name, build_func in tab_configs:
            frame = ttk.Frame(self.notebook, padding=15)
            self.tabs[tab_id] = frame
            self.notebook.add(frame, text=tab_name)
            build_func(frame)

        self.validation_items = {
            "project": ttk.Label(self.tabs["project"], text="❌"),
            "tor": ttk.Label(self.tabs["tor"], text="❌"),
            "apis": ttk.Label(self.tabs["config"], text="❌"),
        }
        self.validation_items["project"].pack(anchor="e")
        self.validation_items["tor"].pack(anchor="e")
        self.validation_items["apis"].pack(anchor="e")

    def _setup_event_handlers(self):
        self.notebook.bind("<<NotebookTabChanged>>", self._on_tab_changed)
        self.master.protocol("WM_DELETE_WINDOW", self._on_window_close)

    def _build_project_tab(self, parent):
        container = ttk.Frame(parent)
        container.pack(fill="both", expand=True)

        self.project_status_label = ttk.Label(container, text="📝 Complete la información del proyecto")
        self.project_status_label.pack(anchor="w", pady=(0, 10))

        form = ttk.Frame(container)
        form.pack(fill="x")

        self.project_vars = {
            "title": tk.StringVar(),
            "country": tk.StringVar(value="Guatemala"),
            "donor": tk.StringVar(),
            "duration_months": tk.StringVar(value="24"),
            "language": tk.StringVar(value="es"),
            "currency": tk.StringVar(value="USD"),
            "municipality": tk.StringVar(),
            "department": tk.StringVar(),
            "beneficiaries_direct": tk.StringVar()
        }
        self.location_vars = {"coverage_type": tk.StringVar(value="Local")}
        self.population_vars = {}

        ttk.Label(form, text="Título").grid(row=0, column=0, sticky="w")
        ttk.Entry(form, textvariable=self.project_vars["title"]).grid(row=0, column=1, sticky="ew")
        ttk.Label(form, text="Donante").grid(row=1, column=0, sticky="w")
        ttk.Entry(form, textvariable=self.project_vars["donor"]).grid(row=1, column=1, sticky="ew")
        ttk.Label(form, text="Duración (meses)").grid(row=2, column=0, sticky="w")
        ttk.Entry(form, textvariable=self.project_vars["duration_months"]).grid(row=2, column=1, sticky="ew")
        ttk.Label(form, text="Moneda").grid(row=3, column=0, sticky="w")
        ttk.Combobox(form, values=["USD", "GTQ", "EUR", "CAD"], textvariable=self.project_vars["currency"]).grid(row=3, column=1, sticky="ew")
        form.columnconfigure(1, weight=1)

        self.target_population_text = tk.Text(container, height=4)
        self.target_population_text.pack(fill="both", expand=False, pady=(10, 0))
        self.org_profile_text = tk.Text(container, height=4)
        self.org_profile_text.pack(fill="both", expand=False, pady=(10, 10))

        buttons = ttk.Frame(container)
        buttons.pack(fill="x")
        ttk.Button(buttons, text="Guardar", command=self._validate_and_save_project).pack(side="left")
        ttk.Button(buttons, text="Limpiar", command=self._clear_project_form).pack(side="left", padx=5)

    def _build_tor_tab(self, parent):
        container = ttk.Frame(parent)
        container.pack(fill="both", expand=True)
        file_row = ttk.Frame(container)
        file_row.pack(fill="x")

        def pick_file():
            p = filedialog.askopenfilename(filetypes=[("Docs", "*.pdf *.docx *.txt"), ("All", "*.*")])
            if p:
                self._on_tor_file_selected(p)

        ttk.Button(file_row, text="Seleccionar ToR", command=pick_file).pack(side="left")
        ttk.Button(file_row, text="Reprocesar", command=self._reprocess_tor).pack(side="left", padx=5)
        ttk.Button(file_row, text="Limpiar", command=self._clear_tor_data).pack(side="left", padx=5)

        self.tor_status_label = ttk.Label(container, text="📄 Seleccione y procese el archivo de Términos de Referencia")
        self.tor_status_label.pack(fill="x", pady=(10, 5))
        self.tor_progress_label = ttk.Label(container, text="")
        self.tor_progress_label.pack(fill="x")

        self.tor_preview_text = tk.Text(container, height=10, state="disabled")
        self.tor_preview_text.pack(fill="both", expand=True, pady=(10, 5))
        self.doc_analysis_text = tk.Text(container, height=4, state="disabled")
        self.doc_analysis_text.pack(fill="x")

    def _build_config_tab(self, parent):
        container = ttk.Frame(parent)
        container.pack(fill="both", expand=True)

        self.api_status_labels = {
            "deepseek": ttk.Label(container, text="❓ Verificando..."),
            "sonnet": ttk.Label(container, text="❓ Verificando...")
        }
        ttk.Label(container, text="Estado DeepSeek:").pack(anchor="w", pady=(0, 2))
        self.api_status_labels["deepseek"].pack(anchor="w")
        ttk.Label(container, text="Estado Sonnet:").pack(anchor="w", pady=(8, 2))
        self.api_status_labels["sonnet"].pack(anchor="w")

        self.model_vars = {
            "narrative_model": tk.StringVar(value="DeepSeek"),
            "budget_model": tk.StringVar(value="Sonnet"),
            "temperature": tk.DoubleVar(value=0.2),
            "max_tokens": tk.IntVar(value=4000),
            "currency": tk.StringVar(value=self.state.get("models.currency", "USD"))
        }
        model_frame = ttk.LabelFrame(container, text="Modelos")
        model_frame.pack(fill="x", pady=(12, 0))
        ttk.Label(model_frame, text="Narrativa").grid(row=0, column=0, sticky="w")
        ttk.Entry(model_frame, textvariable=self.model_vars["narrative_model"]).grid(row=0, column=1, sticky="ew")
        ttk.Label(model_frame, text="Presupuesto").grid(row=1, column=0, sticky="w")
        ttk.Entry(model_frame, textvariable=self.model_vars["budget_model"]).grid(row=1, column=1, sticky="ew")
        ttk.Label(model_frame, text="Temperatura").grid(row=2, column=0, sticky="w")
        ttk.Scale(model_frame, from_=0.0, to=1.0, variable=self.model_vars["temperature"], command=self._update_temp_display).grid(row=2, column=1, sticky="ew")
        ttk.Label(model_frame, text="Max Tokens").grid(row=3, column=0, sticky="w")
        ttk.Entry(model_frame, textvariable=self.model_vars["max_tokens"]).grid(row=3, column=1, sticky="ew")
        ttk.Label(model_frame, text="Moneda").grid(row=4, column=0, sticky="w")
        ttk.Combobox(model_frame, values=["USD", "GTQ", "EUR", "CAD"], textvariable=self.model_vars["currency"]).grid(row=4, column=1, sticky="ew")
        model_frame.columnconfigure(1, weight=1)
        self.temp_display = ttk.Label(model_frame, text=f"{self.model_vars['temperature'].get():.1f}")
        self.temp_display.grid(row=2, column=2, padx=6)

        ttk.Button(container, text="Guardar Configuración", command=self._save_configuration).pack(anchor="e", pady=(10, 0))

    def _build_generate_tab(self, parent):
        container = ttk.Frame(parent)
        container.pack(fill="both", expand=True)

        bar = ttk.Frame(container)
        bar.pack(fill="x")
        self.generate_button = ttk.Button(bar, text="Generar Propuesta", command=self._start_enhanced_generation, state="disabled")
        self.generate_button.pack(side="left")
        self.abort_button = ttk.Button(bar, text="Abortar", command=self._abort_generation, state="disabled")
        self.abort_button.pack(side="left", padx=5)
        self.generate_status_label = ttk.Label(bar, text="❌ Faltan requisitos para la generación")
        self.generate_status_label.pack(side="left", padx=10)

        self.progress_bar = ttk.Progressbar(container, maximum=100)
        self.progress_bar.pack(fill="x", pady=(10, 5))
        self.progress_label = ttk.Label(container, text="")
        self.progress_label.pack(fill="x")

        self.log_text = tk.Text(container, height=18, state="disabled")
        self.log_text.pack(fill="both", expand=True, pady=(10, 0))

    def _build_results_tab(self, parent):
        container = ttk.Frame(parent)
        container.pack(fill="both", expand=True)

        self.results_summary_text = tk.Text(container, height=8, state="disabled")
        self.results_summary_text.pack(fill="x")

        files_row = ttk.Frame(container)
        files_row.pack(fill="both", expand=True)
        self.files_listbox = tk.Listbox(files_row, height=8)
        self.files_listbox.pack(side="left", fill="both", expand=True)
        btns = ttk.Frame(files_row)
        btns.pack(side="left", padx=10)
        ttk.Button(btns, text="Abrir carpeta", command=self._open_results_folder).pack(fill="x", pady=(0, 5))
        ttk.Button(btns, text="Abrir narrativa", command=lambda: self._open_result_file('narrative')).pack(fill="x", pady=(0, 5))
        ttk.Button(btns, text="Abrir presupuesto", command=lambda: self._open_result_file('budget')).pack(fill="x", pady=(0, 5))
        if GANTT_AVAILABLE:
            ttk.Button(btns, text="Abrir Gantt", command=lambda: self._open_result_file('gantt')).pack(fill="x")

    def _start_enhanced_generation(self):
        self._clear_generation_log()
        can_generate = self._run_pre_generation_validation()
        if not can_generate:
            messagebox.showerror("Error", "No se puede iniciar la generación. Verifique los requisitos.")
            return
        self._toggle_processing_state(True, "generation")
        steps = [
            "Inicializando sistemas mejorados...",
            "Obteniendo datos en tiempo real...",
            "Generando narrativa con estándares internacionales...",
            "Creando presupuesto profesional...",
            "Generando cronograma Gantt...",
            "Guardando archivos y documentos..."
        ]
        self.progress_manager.set_steps(steps)
        self._current_thread = threading.Thread(target=self._enhanced_generation_task, daemon=True)
        self._current_thread.start()

    def _enhanced_generation_task(self):
        try:
            project_info = self.state.get("project", {})
            tor_content = self.state.get("tor_content", "")
            currency = self.state.get("models.currency", "USD")

            self.progress_manager.next_step()
            self._log_message("Inicializando sistemas mejorados...")
            self._initialize_enhanced_systems(project_info, tor_content)

            self.progress_manager.next_step()
            self._log_message("Obteniendo indicadores económicos actuales...")
            realtime_data = None
            if self.realtime_system and REALTIME_DATA_AVAILABLE:
                try:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    realtime_data = loop.run_until_complete(self.realtime_system.fetch_current_indicators())
                    loop.close()
                    self._log_message("Datos económicos actualizados obtenidos")
                except Exception as e:
                    self._log_message(f"Advertencia: No se pudieron obtener datos en tiempo real: {e}")

            self.progress_manager.next_step()
            self._log_message("Generando narrativa con estándares internacionales...")
            narrative_result = self._generate_enhanced_narrative(project_info, tor_content, realtime_data)
            self.state.set("results.narrative", narrative_result)
            self._log_message("Narrativa profesional generada exitosamente")

            self.progress_manager.next_step()
            self._log_message("Creando presupuesto profesional...")
            budget_result = self._generate_enhanced_budget(project_info, tor_content, currency)
            self.state.set("results.budget", json.dumps(budget_result, ensure_ascii=False))
            self._log_message("Presupuesto profesional generado exitosamente")

            gantt_data = None
            if GANTT_AVAILABLE:
                self.progress_manager.next_step()
                self._log_message("Generando cronograma Gantt profesional...")
                try:
                    gantt_data = self._generate_gantt_chart(project_info, budget_result)
                    self._log_message("Cronograma Gantt generado exitosamente")
                except Exception as e:
                    self._log_message(f"Advertencia: Error generando Gantt: {e}")

            self.progress_manager.next_step()
            self._log_message("Guardando todos los archivos...")
            output_folder = self._save_enhanced_results(narrative_result, budget_result, gantt_data, currency)

            self.state.set("results.output_paths", {
                "folder": str(output_folder),
                "narrative_docx": str(output_folder / "propuesta_narrativa.docx"),
                "budget_xlsx": str(output_folder / "presupuesto_detallado.xlsx"),
                "gantt_xlsx": str(output_folder / "cronograma_gantt.xlsx") if gantt_data else None
            })

            self._log_message("Todos los archivos guardados exitosamente")
            self.progress_manager.update("Generación completa", 100)
            self.master.after(0, self._finish_generation_ui_update, True)

        except Exception as e:
            traceback.print_exc()
            error_msg = f"Error crítico en generación mejorada: {str(e)}"
            self._log_message(error_msg)
            self.master.after(0, lambda: messagebox.showerror("Error de Generación", error_msg))
            self.master.after(0, self._finish_generation_ui_update, False)

    def _initialize_enhanced_systems(self, project_info: Dict, tor_content: str):
        if ENHANCED_BUDGET_AVAILABLE:
            self.budget_generator = InternationalBudgetGenerator(project_info, tor_content)
            self._log_message("Sistema de presupuesto profesional inicializado")
        if ACADEMIC_REFS_AVAILABLE:
            self.academic_system = AcademicReferenceSystem(project_info, tor_content)
            self._log_message("Sistema de referencias académicas inicializado")
        if REALTIME_DATA_AVAILABLE:
            self.realtime_system = RealTimeDataIntegrator(project_info)
            self._log_message("Sistema de datos en tiempo real inicializado")
        if GANTT_AVAILABLE:
            self._log_message("Sistema de cronograma Gantt disponible")

    def _generate_enhanced_narrative(self, project_info: Dict, tor_content: str, realtime_data: Dict = None) -> str:
        try:
            deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
            if not deepseek_api_key:
                return "Error: API key de DeepSeek no configurada"
            deepseek_client = DeepSeekClient(deepseek_api_key)
            generator = EnhancedChainedPromptGenerator(deepseek_client, max_tokens_per_chunk=8000, progress_callback=self._log_message)
            tor_chunks = [{"content": tor_content, "section": "complete"}]
            if len(tor_content or "") > 10000:
                tor_chunks = TokenManager.intelligent_chunk_tor(tor_content, 8000)
            result = generator.process_tor_chunks(tor_chunks, project_info, "narrative")
            if isinstance(result, str) and not result.startswith("Error"):
                return result
            return self._generate_basic_narrative(project_info, tor_content)
        except Exception as e:
            self._log_message(f"Error en generación de narrativa: {e}")
            return self._generate_basic_narrative(project_info, tor_content)

    def _generate_enhanced_budget(self, project_info: Dict, tor_content: str, currency: str) -> Dict:
        try:
            if ENHANCED_BUDGET_AVAILABLE and self.budget_generator:
                self._log_message("Usando generador de presupuesto internacional...")
                budget_result = self.budget_generator.generate_comprehensive_budget()
                budget_result["currency"] = currency
                self._log_message(f"Presupuesto generado en {currency} con {len(budget_result.get('items', []))} líneas")
                return budget_result
            else:
                return self._generate_sonnet_budget(project_info, tor_content, currency)
        except Exception as e:
            self._log_message(f"Error en generación de presupuesto: {e}")
            return self._generate_fallback_budget(project_info, currency)

    def _generate_sonnet_budget(self, project_info: Dict, tor_content: str, currency: str) -> Dict:
        try:
            sonnet_api_key = os.getenv("SONNET_API_KEY") or os.getenv("ANTHROPIC_API_KEY")
            if not sonnet_api_key:
                raise Exception("API key de Sonnet no configurada")
            sonnet_client = SonnetClient(sonnet_api_key)
            budget_prompt = self._build_comprehensive_budget_prompt(project_info, tor_content, currency)
            budget_schema = self._get_comprehensive_budget_schema(currency)
            budget_result = sonnet_client.generate_json(budget_prompt, budget_schema)
            if budget_result and not budget_result.get("error"):
                self._log_message(f"Presupuesto Sonnet generado: {len(budget_result.get('items', []))} líneas")
                return budget_result
            else:
                raise Exception(budget_result.get("error", "Error desconocido"))
        except Exception as e:
            self._log_message(f"Error en Sonnet: {e}")
            return self._generate_fallback_budget(project_info, currency)

    def _build_comprehensive_budget_prompt(self, project_info: Dict, tor_content: str, currency: str) -> str:
        return f"""
Eres un analista financiero senior especializado en proyectos de desarrollo internacional. 

INFORMACIÓN DEL PROYECTO:
- Título: {project_info.get('title', 'N/A')}
- Donante: {project_info.get('donor', 'Internacional')}
- Duración: {project_info.get('duration_months', 24)} meses
- Moneda: {currency}
- Ubicación: {project_info.get('municipality', '')}, {project_info.get('department', '')}, {project_info.get('country', 'Guatemala')}
- Beneficiarios directos: {project_info.get('beneficiaries_direct', 'Por definir')}

TÉRMINOS DE REFERENCIA:
{(tor_content or '')[:8000]}

GENERA un presupuesto profesional en formato JSON con las siguientes características:

1. MONEDA: Todas las cantidades deben estar en {currency}
2. CATEGORÍAS: Usa categorías estándar internacionales
3. DETALLE: Mínimo 25-40 líneas presupuestales específicas
4. CÁLCULOS: Asegura que total_cost = quantity * unit_cost
5. REALISMO: Costos realistas para Guatemala/{currency}

ESTRUCTURA JSON REQUERIDA:
{{
  "currency": "{currency}",
  "exchange_rate": {7.75 if currency == "GTQ" else 1.0},
  "project_duration_months": {project_info.get('duration_months', 24)},
  "items": [
    {{
      "code": "1.1.1",
      "category": "Personal Técnico y Profesional",
      "description": "Descripción específica",
      "unit": "mes",
      "quantity": 24,
      "unit_cost": 2500,
      "total_cost": 60000,
      "justification": "Justificación técnica"
    }}
  ],
  "summary_by_category": {{"Personal Técnico y Profesional": 150000}},
  "financial_totals": {{
    "subtotal": 200000,
    "administrative_cost": 20000,
    "contingency_amount": 15000,
    "total": 235000
  }},
  "assumptions": ["Lista de supuestos presupuestales"],
  "compliance_notes": ["Cumplimiento normativo"]
}}

DEVUELVE SOLO EL JSON, SIN EXPLICACIONES ADICIONALES.
"""

    def _get_comprehensive_budget_schema(self, currency: str) -> Dict:
        return {
            "currency": currency,
            "exchange_rate": "number",
            "project_duration_months": "integer",
            "items": [
                {
                    "code": "string",
                    "category": "string",
                    "description": "string",
                    "unit": "string",
                    "quantity": "number",
                    "unit_cost": "number",
                    "total_cost": "number",
                    "justification": "string"
                }
            ],
            "summary_by_category": {"string": "number"},
            "financial_totals": {
                "subtotal": "number",
                "administrative_cost": "number",
                "contingency_amount": "number",
                "total": "number"
            },
            "assumptions": ["string"],
            "compliance_notes": ["string"]
        }

    def _generate_fallback_budget(self, project_info: Dict, currency: str) -> Dict:
        duration = int(project_info.get('duration_months', 24) or 24)
        if currency == "GTQ":
            rate = 7.75
        elif currency == "EUR":
            rate = 0.85
        elif currency == "CAD":
            rate = 1.35
        else:
            rate = 1.0
        base_costs = {
            "coordinator": 2800 * rate,
            "specialist": 2200 * rate,
            "facilitator": 1000 * rate,
        }
        items = [
            {
                "code": "1.1.1",
                "category": "Personal Técnico y Profesional",
                "description": "Coordinador de Proyecto Senior",
                "unit": "mes",
                "quantity": duration,
                "unit_cost": base_costs["coordinator"],
                "total_cost": duration * base_costs["coordinator"],
                "justification": "Coordinación general del proyecto"
            },
            {
                "code": "1.1.2",
                "category": "Personal Técnico y Profesional",
                "description": "Especialista Técnico",
                "unit": "mes",
                "quantity": duration,
                "unit_cost": base_costs["specialist"],
                "total_cost": duration * base_costs["specialist"],
                "justification": "Asesoría técnica especializada"
            }
        ]
        subtotal = sum(i["total_cost"] for i in items)
        admin_cost = subtotal * 0.08
        contingency = subtotal * 0.05
        total = subtotal + admin_cost + contingency
        return {
            "currency": currency,
            "exchange_rate": rate,
            "project_duration_months": duration,
            "items": items,
            "summary_by_category": {
                "Personal Técnico y Profesional": sum(i["total_cost"] for i in items)
            },
            "financial_totals": {
                "subtotal": subtotal,
                "administrative_cost": admin_cost,
                "contingency_amount": contingency,
                "total": total
            },
            "assumptions": [f"Tipo de cambio {currency}/USD: {rate}"],
            "compliance_notes": ["Presupuesto generado con estándares internacionales"]
        }

    def _generate_basic_narrative(self, project_info: Dict, tor_content: str) -> str:
        return f"""
PROPUESTA DE PROYECTO: {project_info.get('title', 'Proyecto de Desarrollo')}

1. RESUMEN EJECUTIVO
El presente proyecto propone una intervención integral de {project_info.get('duration_months', '24')} meses en {project_info.get('municipality', '')}, {project_info.get('department', '')}, Guatemala.

2. CONTEXTO Y JUSTIFICACIÓN
Basándose en los términos de referencia proporcionados, se ha identificado la necesidad de una intervención específica.

3. OBJETIVOS
Objetivo General: Contribuir al desarrollo sostenible de las comunidades beneficiarias.

4. METODOLOGÍA
IEPADES aplicará su experiencia utilizando metodologías participativas.

5. CRONOGRAMA
El proyecto se ejecutará durante {project_info.get('duration_months', '24')} meses.

6. PRESUPUESTO
El presupuesto se ajusta a los parámetros establecidos.
"""

    def _generate_gantt_chart(self, project_info: Dict, budget_data: Dict) -> Dict:
        if not GANTT_AVAILABLE:
            return None
        try:
            activities = self._extract_activities_from_budget(budget_data)
            self.gantt_generator = GanttChartGenerator(project_info, activities, budget_data)
            return {"activities": activities, "project_info": project_info, "budget_data": budget_data}
        except Exception as e:
            self._log_message(f"Error generando Gantt: {e}")
            return None

    def _on_tab_changed(self, event):
        current_tab = self.notebook.tab(self.notebook.select(), "text")
        if current_tab == "🚀 Generación":
            self._run_pre_generation_validation()

    def _update_progress_ui(self, progress: int, step_text: str):
        self.master.after(0, self._set_progress_state, progress, step_text)

    def _set_progress_state(self, progress: int, step_text: str):
        self.progress_bar['value'] = progress
        self.progress_label['text'] = step_text
        self.state.set("processing.current_step", step_text)
        self.state.set("processing.progress", progress)

    def _validate_and_save_project(self):
        self._update_state_from_project_ui()
        is_valid, errors = self.state.validate_project()
        if is_valid:
            messagebox.showinfo("Éxito", "Datos del proyecto guardados y validados correctamente.")
            self._update_project_status_ui(True)
        else:
            messagebox.showerror("Error de Validación", "\n".join(errors))
            self._update_project_status_ui(False)

    def _update_state_from_project_ui(self):
        project_data = {}
        for key, var in self.project_vars.items():
            project_data[key] = var.get()
        for key, var in self.location_vars.items():
            project_data[key] = var.get()
        for key, var in self.population_vars.items():
            project_data[key] = var.get()
        project_data['target_population'] = self.target_population_text.get("1.0", tk.END).strip()
        project_data['org_profile'] = self.org_profile_text.get("1.0", tk.END).strip()
        self.state.set("project", project_data)

    def _update_project_status_ui(self, is_valid: bool):
        if is_valid:
            self.project_status_label['text'] = "✅ Datos del proyecto validados y listos"
            self.project_status_label['foreground'] = "green"
        else:
            self.project_status_label['text'] = "❌ Hay errores en los datos del proyecto"
            self.project_status_label['foreground'] = "red"

    def _clear_project_form(self):
        for var in self.project_vars.values():
            var.set("")
        for var in self.location_vars.values():
            var.set("")
        for var in self.population_vars.values():
            var.set("")
        self.target_population_text.delete("1.0", tk.END)
        self.org_profile_text.delete("1.0", tk.END)
        self.project_vars['country'].set("Guatemala")
        self.project_vars['language'].set("es")
        self.project_vars['currency'].set("USD")
        self.location_vars['coverage_type'].set("Local")
        self._update_project_status_ui(False)

    def _on_tor_file_selected(self, file_path):
        self.state.set("tor_path", file_path)
        self._start_tor_processing(file_path)

    def _reprocess_tor(self):
        tor_path = self.state.get("tor_path")
        if tor_path:
            self._start_tor_processing(tor_path)
        else:
            messagebox.showerror("Error", "No hay un archivo de ToR seleccionado para reprocesar.")

    def _clear_tor_data(self):
        self.state.set("tor_path", None)
        self.state.set("tor_content", None)
        self.state.set("tor_chunks", {})
        try:
            self.tor_progress_label.config(text="")
            self.tor_preview_text.config(state="normal")
            self.tor_preview_text.delete("1.0", tk.END)
            self.tor_preview_text.config(state="disabled")
            self.doc_analysis_text.config(state="normal")
            self.doc_analysis_text.delete("1.0", tk.END)
            self.doc_analysis_text.config(state="disabled")
            self.validation_items["tor"].config(text="❌", foreground="red")
        except:
            pass

    def _start_tor_processing(self, file_path):
        self._toggle_processing_state(True, "tor")
        self._current_thread = threading.Thread(target=self._process_tor_task, args=(file_path,), daemon=True)
        self._current_thread.start()

    def _process_tor_task(self, file_path):
        try:
            self.master.after(0, lambda: self.tor_progress_label.config(text="Procesando documento..."))
            processor = DocumentProcessor()
            tor_content = processor.extract_text_from_file(file_path)
            if not tor_content or (isinstance(tor_content, str) and tor_content.startswith("Error")):
                raise Exception("No se pudo extraer contenido del archivo")
            self.state.set("tor_content", tor_content)
            self.master.after(0, lambda: self.tor_progress_label.config(text="Analizando contenido..."))
            tor_chunks = processor.chunk_document(tor_content)
            self.state.set("tor_chunks", tor_chunks)
            analysis_text = f"Análisis: {len(tor_content)} caracteres, {tor_chunks.get('section_count', 0)} secciones"
            self.master.after(0, self._safe_update_tor_ui, tor_content, analysis_text, True)
            self.master.after(0, lambda: self._toggle_processing_state(False, "tor"))
        except Exception as e:
            self.master.after(0, self._safe_handle_tor_error, str(e))

    def _safe_update_tor_ui(self, content, analysis, success):
        try:
            self.tor_progress_label.config(text="Procesamiento completado" if success else "Error al procesar")
            self.tor_preview_text.config(state="normal")
            self.tor_preview_text.delete("1.0", tk.END)
            if content:
                self.tor_preview_text.insert("1.0", content[:20000])
            self.tor_preview_text.config(state="disabled")
            self.doc_analysis_text.config(state="normal")
            self.doc_analysis_text.delete("1.0", tk.END)
            if analysis:
                self.doc_analysis_text.insert("1.0", analysis)
            self.doc_analysis_text.config(state="disabled")
            self.validation_items["tor"].config(text="✅" if success else "❌", foreground="green" if success else "red")
        except Exception as e:
            self._log_message(f"Error actualizando UI de ToR: {e}")
            try:
                self.validation_items["tor"].config(text="❌", foreground="red")
            except:
                pass

    def _safe_handle_tor_error(self, error_msg: str):
        try:
            self.tor_progress_label.config(text=f"Error: {error_msg}")
            self.tor_preview_text.config(state="normal")
            self.tor_preview_text.delete("1.0", tk.END)
            self.tor_preview_text.insert("1.0", f"Error procesando ToR: {error_msg}")
            self.tor_preview_text.config(state="disabled")
            self.doc_analysis_text.config(state="normal")
            self.doc_analysis_text.delete("1.0", tk.END)
            self.doc_analysis_text.insert("1.0", "No se pudo analizar el documento.")
            self.doc_analysis_text.config(state="disabled")
            self.validation_items["tor"].config(text="❌", foreground="red")
            self._toggle_processing_state(False, "tor")
        except:
            pass

    def _check_api_status_async(self):
        self.api_status_labels['deepseek']['text'] = "❓ Verificando..."
        self.api_status_labels['sonnet']['text'] = "❓ Verificando..."
        threading.Thread(target=self._check_api_status_task, daemon=True).start()

    def _check_api_status_task(self):
        try:
            deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
            sonnet_api_key = os.getenv("SONNET_API_KEY") or os.getenv("ANTHROPIC_API_KEY")

            deepseek_status = bool(deepseek_api_key)
            sonnet_status = bool(sonnet_api_key)

            if deepseek_api_key:
                try:
                    deepseek_client = DeepSeekClient(deepseek_api_key)
                    result = deepseek_client.generate("test")
                    deepseek_status = bool(getattr(result, "content", "") or result)
                except:
                    deepseek_status = False

            if sonnet_api_key:
                try:
                    SonnetClient(sonnet_api_key)
                    sonnet_status = True
                except:
                    sonnet_status = False

            self.master.after(0, self._update_api_status_ui, deepseek_status, sonnet_status)
        except Exception:
            self.master.after(0, self._update_api_status_ui, False, False)

    def _update_api_status_ui(self, deepseek_ok: bool, sonnet_ok: bool):
        self.state.set("api_status.deepseek", deepseek_ok)
        self.state.set("api_status.sonnet", sonnet_ok)
        self.api_status_labels['deepseek']['text'] = "✅ Conectado" if deepseek_ok else "❌ Desconectado"
        self.api_status_labels['deepseek']['foreground'] = "green" if deepseek_ok else "red"
        self.api_status_labels['sonnet']['text'] = "✅ Conectado" if sonnet_ok else "❌ Desconectado"
        self.api_status_labels['sonnet']['foreground'] = "green" if sonnet_ok else "red"
        self._run_pre_generation_validation()

    def _update_temp_display(self, *args):
        temp = self.model_vars['temperature'].get()
        self.temp_display['text'] = f"{temp:.1f}"

    def _save_configuration(self):
        models_config = {
            'narrative': self.model_vars['narrative_model'].get(),
            'budget': self.model_vars['budget_model'].get(),
            'temperature': self.model_vars['temperature'].get(),
            'max_tokens': self.model_vars['max_tokens'].get(),
            'currency': self.model_vars['currency'].get()
        }
        self.state.set("models", models_config)
        messagebox.showinfo("Éxito", "Configuración guardada correctamente.")

    def _run_pre_generation_validation(self):
        project_ok, _ = self.state.validate_project()
        tor_ok, _ = self.state.validate_tor()
        apis_ok, _ = self.state.validate_apis()
        self._update_validation_ui("project", project_ok)
        self._update_validation_ui("tor", tor_ok)
        self._update_validation_ui("apis", apis_ok)
        can_generate = project_ok and tor_ok and apis_ok
        if can_generate:
            try:
                self.generate_button.config(state="normal")
                self.generate_status_label.config(text="🚀 Listo para generar la propuesta", foreground="green")
            except Exception:
                pass
        else:
            try:
                self.generate_button.config(state="disabled")
            except Exception:
                pass
            self.generate_status_label.config(text="❌ Faltan requisitos para la generación", foreground="red")
        return can_generate

    def _update_validation_ui(self, item_id, is_ok):
        if item_id in self.validation_items:
            self.validation_items[item_id]['text'] = "✅" if is_ok else "❌"
            self.validation_items[item_id]['foreground'] = "green" if is_ok else "red"

    def _clear_generation_log(self):
        self.log_text.config(state="normal")
        self.log_text.delete("1.0", tk.END)
        self.log_text.config(state="disabled")

    def _finish_generation_ui_update(self, success):
        self._toggle_processing_state(False, "generation")
        if success:
            messagebox.showinfo("Generación Completa", "La propuesta ha sido generada con éxito.")
        self._update_results_tab_ui()

    def _abort_generation(self):
        if self._current_thread and self._current_thread.is_alive():
            messagebox.showwarning("Abortar", "La generación se detendrá.")
            self._toggle_processing_state(False, "generation")

    def _toggle_processing_state(self, is_processing, tab):
        self._processing = is_processing
        if tab == "generation":
            try:
                self.generate_button.config(state="disabled" if is_processing else ("normal" if self._run_pre_generation_validation() else "disabled"))
            except Exception:
                pass
            try:
                self.abort_button.config(state="normal" if is_processing else "disabled")
            except Exception:
                pass
        if tab == "tor":
            pass

    def _update_results_tab_ui(self):
        narrative_path = self.state.get("results.output_paths.narrative_docx")
        budget_path = self.state.get("results.output_paths.budget_xlsx")
        summary_text = ""
        files_list = []
        if narrative_path:
            summary_text += "Documento de Narrativa: ✅ Generado\n"
            files_list.append(f"Narrativa: {Path(narrative_path).name}")
        if budget_path:
            summary_text += "Documento de Presupuesto: ✅ Generado\n"
            files_list.append(f"Presupuesto: {Path(budget_path).name}")
        self.results_summary_text.config(state="normal")
        self.results_summary_text.delete("1.0", tk.END)
        self.results_summary_text.insert("1.0", summary_text or "No hay resultados generados.")
        self.results_summary_text.config(state="disabled")
        self.files_listbox.delete(0, tk.END)
        for item in files_list:
            self.files_listbox.insert(tk.END, item)

    def _open_results_folder(self):
        folder_path = self.state.get("results.output_paths.folder")
        if folder_path and os.path.exists(folder_path):
            try:
                if os.name == 'nt':
                    os.startfile(folder_path)
                elif os.name == 'posix':
                    import subprocess
                    if sys.platform == 'darwin':
                        subprocess.Popen(['open', folder_path])
                    else:
                        subprocess.Popen(['xdg-open', folder_path])
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo abrir la carpeta: {e}")
        else:
            messagebox.showerror("Error", "La carpeta de resultados no existe.")

    def _open_result_file(self, file_type):
        if file_type == 'narrative':
            path = self.state.get("results.output_paths.narrative_docx")
        elif file_type == 'budget':
            path = self.state.get("results.output_paths.budget_xlsx")
        elif file_type == 'gantt':
            path = self.state.get("results.output_paths.gantt_xlsx")
        else:
            path = None
        if path and os.path.exists(path):
            try:
                if os.name == 'nt':
                    os.startfile(path)
                elif os.name == 'posix':
                    import subprocess
                    if sys.platform == 'darwin':
                        subprocess.Popen(['open', path])
                    else:
                        subprocess.Popen(['xdg-open', path])
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo abrir el archivo: {e}")
        else:
            messagebox.showerror("Error", "El archivo no existe. Genere la propuesta primero.")

    def _on_window_close(self):
        if self._processing:
            if messagebox.askyesno("Confirmar Salida", "Hay un proceso en curso. ¿Desea salir?"):
                self.master.destroy()
        else:
            self.master.destroy()


def main():
    root = tk.Tk()
    root.title("Asistente para Propuestas de Proyectos - IEPADES Enhanced")
    root.geometry("900x700")
    root.minsize(800, 600)
    style = ttk.Style(root)
    try:
        available_themes = style.theme_names()
        if 'vista' in available_themes:
            style.theme_use('vista')
        elif 'clam' in available_themes:
            style.theme_use('clam')
        else:
            style.theme_use('default')
    except Exception:
        style.theme_use('default')
    style.configure('Enhanced.TNotebook', tabposition='n')
    style.configure('Enhanced.TNotebook.Tab', padding=[12, 8])
    try:
        app = ProposalWizard(root)
        root.update_idletasks()
        x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
        y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
        root.geometry(f"+{x}+{y}")
        root.deiconify()
        root.mainloop()
    except KeyboardInterrupt:
        print("\nAplicación interrumpida por el usuario.")
    except Exception as e:
        print(f"Error crítico en la aplicación: {e}")
        traceback.print_exc()
        messagebox.showerror("Error Crítico", f"Ocurrió un error crítico:\n{str(e)}")
    finally:
        try:
            root.destroy()
        except:
            pass


if __name__ == "__main__":
    main()
